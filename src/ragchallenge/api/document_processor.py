"""
Document Processing Module
Handles document upload, processing, and vector store management for RAG system.
"""

import os
import tempfile
import uuid
from pathlib import Path
from typing import List, Optional
import aiofiles
from fastapi import UploadFile, HTTPException

# Document processing imports
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings

from .config import Settings


class DocumentProcessor:
    """Handles document upload, processing, and vector store management."""
    
    def __init__(self, config: Settings):
        self.config = config
        self.upload_dir = Path("data/uploads")
        self.upload_dir.mkdir(exist_ok=True)
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Lazy-load embeddings
        self.embeddings = None
    
    def get_embeddings(self):
        """Lazy-load embeddings when needed."""
        if self.embeddings is None:
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.config.embedding_model,
                model_kwargs={'device': self.config.embedding_model_device}
            )
        return self.embeddings
    
    async def save_upload_file(self, upload_file: UploadFile) -> str:
        """Save uploaded file to disk and return file path."""
        # Generate unique filename
        file_extension = Path(upload_file.filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = self.upload_dir / unique_filename
        
        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            content = await upload_file.read()
            await f.write(content)
        
        return str(file_path)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error processing TXT: {str(e)}")
    
    def extract_text_from_file(self, file_path: str, filename: str) -> str:
        """Extract text from file based on its extension."""
        file_extension = Path(filename).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            return self.extract_text_from_txt(file_path)
        else:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {file_extension}. Supported types: PDF, DOCX, TXT, MD"
            )
    
    def create_documents_from_text(self, text: str, filename: str) -> List[LangchainDocument]:
        """Split text into chunks and create LangChain documents."""
        chunks = self.text_splitter.split_text(text)
        documents = []
        
        for i, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": filename,
                    "chunk_id": i,
                    "document_type": "uploaded_document"
                }
            )
            documents.append(doc)
        
        return documents
    
    def create_user_vectorstore(self, user_id: str) -> str:
        """Create a user-specific vector store directory."""
        user_vectorstore_dir = Path(f"data/user_vectorstores/{user_id}")
        user_vectorstore_dir.mkdir(parents=True, exist_ok=True)
        return str(user_vectorstore_dir)
    
    async def process_and_store_document(self, upload_file: UploadFile, user_id: Optional[str] = None) -> dict:
        """Process uploaded document and add to vector store."""
        try:
            # Validate file type
            allowed_extensions = ['.pdf', '.docx', '.txt', '.md']
            file_extension = Path(upload_file.filename).suffix.lower()
            
            if file_extension not in allowed_extensions:
                raise HTTPException(
                    status_code=400,
                    detail=f"File type {file_extension} not supported. Allowed types: {', '.join(allowed_extensions)}"
                )
            
            # Save uploaded file
            file_path = await self.save_upload_file(upload_file)
            
            try:
                # Extract text from file
                text = self.extract_text_from_file(file_path, upload_file.filename)
                
                if not text.strip():
                    raise HTTPException(status_code=400, detail="No text content found in the file")
                
                # Create documents
                documents = self.create_documents_from_text(text, upload_file.filename)
                
                # Determine vector store path
                if user_id:
                    vectorstore_path = self.create_user_vectorstore(user_id)
                else:
                    # Use default augmented vectorstore
                    vectorstore_path = self.config.data_dir
                
                # Get embeddings
                embeddings = self.get_embeddings()
                
                # Load existing vectorstore or create new one
                try:
                    vectorstore = Chroma(
                        persist_directory=vectorstore_path,
                        embedding_function=embeddings
                    )
                    
                    # Add documents to vectorstore
                    vectorstore.add_documents(documents)
                    
                    # Persist the vectorstore
                    vectorstore.persist()
                    
                except Exception as vs_error:
                    # If vectorstore doesn't exist, create it
                    vectorstore = Chroma.from_documents(
                        documents,
                        embeddings,
                        persist_directory=vectorstore_path
                    )
                
                return {
                    "status": "success",
                    "message": f"Successfully processed {upload_file.filename}",
                    "document_name": upload_file.filename,
                    "chunks_created": len(documents),
                    "vectorstore_path": vectorstore_path,
                    "text_preview": text[:200] + "..." if len(text) > 200 else text
                }
                
            finally:
                # Clean up uploaded file
                if os.path.exists(file_path):
                    os.remove(file_path)
                    
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
    
    def list_user_documents(self, user_id: str) -> List[dict]:
        """List documents in user's vector store."""
        user_vectorstore_path = Path(f"data/user_vectorstores/{user_id}")
        
        if not user_vectorstore_path.exists():
            return []
        
        try:
            vectorstore = Chroma(
                persist_directory=str(user_vectorstore_path),
                embedding_function=self.get_embeddings()
            )
            
            # Get all documents
            collection = vectorstore._collection
            results = collection.get()
            
            # Extract unique document sources
            documents = {}
            for metadata in results.get('metadatas', []):
                if metadata and 'source' in metadata:
                    source = metadata['source']
                    if source not in documents:
                        documents[source] = {
                            "name": source,
                            "chunks": 0,
                            "document_type": metadata.get('document_type', 'unknown')
                        }
                    documents[source]["chunks"] += 1
            
            return list(documents.values())
            
        except Exception as e:
            return []
    
    def delete_user_document(self, user_id: str, document_name: str) -> dict:
        """Delete a specific document from user's vector store."""
        user_vectorstore_path = Path(f"data/user_vectorstores/{user_id}")
        
        if not user_vectorstore_path.exists():
            raise HTTPException(status_code=404, detail="User vector store not found")
        
        try:
            vectorstore = Chroma(
                persist_directory=str(user_vectorstore_path),
                embedding_function=self.get_embeddings()
            )
            
            # Get collection and delete documents by source
            collection = vectorstore._collection
            
            # Find documents with matching source
            results = collection.get(where={"source": document_name})
            
            if not results.get('ids'):
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Delete the documents
            collection.delete(ids=results['ids'])
            
            return {
                "status": "success",
                "message": f"Successfully deleted {document_name}",
                "deleted_chunks": len(results['ids'])
            }
            
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")